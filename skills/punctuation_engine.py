from __future__ import annotations

import re

from docx import Document


LEFT_DOUBLE_QUOTE = "\u201c"
RIGHT_DOUBLE_QUOTE = "\u201d"
LEFT_SINGLE_QUOTE = "\u2018"
RIGHT_SINGLE_QUOTE = "\u2019"

REPLACEMENTS = {
    "(": "（",
    ")": "）",
    ":": "：",
    ";": "；",
    "?": "？",
    "!": "！",
}

_PLACEHOLDER_PREFIX = "\x02PROT"


def _protect_special_patterns(text: str) -> tuple[str, list[tuple[str, str]]]:
    protected: list[tuple[str, str]] = []
    counter = [0]

    def _replace_with_placeholder(match):
        placeholder = f"{_PLACEHOLDER_PREFIX}{counter[0]}\x03"
        protected.append((placeholder, match.group()))
        counter[0] += 1
        return placeholder

    result = text
    result = re.sub(r"(?:https?|ftp)://\S+", _replace_with_placeholder, result)
    result = re.sub(r"[\w.+-]+@[\w-]+\.[\w.-]+", _replace_with_placeholder, result)
    result = re.sub(r"[A-Za-z]:\\", _replace_with_placeholder, result)
    result = re.sub(r"[A-Za-z]+[\s-]?\d+:\d{2,}", _replace_with_placeholder, result)
    result = re.sub(r"(?<!\d)(\d{1,2}:\d{2}(?::\d{2})?)(?!\d)", _replace_with_placeholder, result)
    return result, protected


def _restore_protected(text: str, protected: list[tuple[str, str]]) -> str:
    result = text
    for placeholder, original in protected:
        result = result.replace(placeholder, original)
    return result


def has_chinese(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def _fix_simple_punctuation(text: str) -> str:
    if not text:
        return text

    result, protected = _protect_special_patterns(text)
    result = re.sub(r"\.{2,}", "……", result)
    result = re.sub(r"。{2,}", "……", result)
    result = re.sub(r"--+", "——", result)
    result = re.sub(r"—(?!—)", "——", result)

    if has_chinese(result):
        for en, cn in REPLACEMENTS.items():
            result = result.replace(en, cn)

    result = re.sub(r"([\u4e00-\u9fff]),", r"\1，", result)
    result = re.sub(r",([\u4e00-\u9fff])", r"，\1", result)
    result = re.sub(r"([\u4e00-\u9fff])\.(\s|$)", r"\1。\2", result)
    return _restore_protected(result, protected)


def _fix_quotes_whole_text(text: str) -> str:
    result = text

    double_quote_chars = ['"', "\u201c", "\u201d", "\u201e", "\u201f", "\u300c", "\u300d"]
    temp = result
    for q in double_quote_chars:
        temp = temp.replace(q, "\x00")

    if "\x00" in temp:
        chars = list(temp)
        quote_idx = 0
        for i, c in enumerate(chars):
            if c == "\x00":
                chars[i] = LEFT_DOUBLE_QUOTE if quote_idx % 2 == 0 else RIGHT_DOUBLE_QUOTE
                quote_idx += 1
        result = "".join(chars)

    single_quote_chars = ["'", "\u2018", "\u2019", "\u201a", "\u201b"]
    temp = result
    for q in single_quote_chars:
        temp = temp.replace(q, "\x01")

    if "\x01" in temp:
        chars = list(temp)
        quote_idx = 0
        for i, c in enumerate(chars):
            if c == "\x01":
                chars[i] = LEFT_SINGLE_QUOTE if quote_idx % 2 == 0 else RIGHT_SINGLE_QUOTE
                quote_idx += 1
        result = "".join(chars)

    return result


def _redistribute_text_to_runs(runs, new_full_text: str) -> None:
    run_lengths = [len(run.text) for run in runs]
    total_original = sum(run_lengths)
    if len(new_full_text) == total_original:
        pos = 0
        for i, run in enumerate(runs):
            run.text = new_full_text[pos : pos + run_lengths[i]]
            pos += run_lengths[i]
    else:
        runs[0].text = new_full_text
        for run in runs[1:]:
            run.text = ""


def process_paragraph(para) -> bool:
    full_text = para.text
    if not full_text.strip():
        return False

    runs = para.runs
    if not runs:
        return False

    changed = False
    for run in runs:
        original = run.text
        fixed = _fix_simple_punctuation(original)
        if fixed != original:
            run.text = fixed
            changed = True

    full_after_simple = para.text
    full_after_quotes = _fix_quotes_whole_text(full_after_simple)

    if full_after_quotes != full_after_simple:
        _redistribute_text_to_runs(runs, full_after_quotes)
        changed = True

    return changed


def process_document(input_path: str, output_path: str) -> None:
    print(f"Reading: {input_path}")
    doc = Document(input_path)

    changes = 0
    for i, para in enumerate(doc.paragraphs):
        if process_paragraph(para):
            changes += 1
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            print(f"  Para {i + 1}: {preview}")

    table_changes = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para):
                        table_changes += 1

    if table_changes > 0:
        print(f"  Tables: {table_changes} cells fixed")

    print()
    print(f"Total: {changes} paragraphs + {table_changes} table cells fixed")
    doc.save(output_path)
    print(f"Saved: {output_path}")
